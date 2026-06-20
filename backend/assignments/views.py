import io
import logging
import re
import asyncio
from rest_framework import viewsets, permissions, status
from rest_framework.response import Response
from rest_framework.decorators import action
from django.shortcuts import get_object_or_404
from django.http import HttpResponse
from django.utils import timezone

from .models import Assignment
from .serializers import AssignmentSerializer

logger = logging.getLogger('nitemind')


# ─────────────────────────────────────────────────────────────────────────────
# LaTeX → Human-Readable Converter
# Applied to ai_response text before any export (PDF, DOCX, TXT).
# Converts LaTeX math syntax into Unicode symbols so exported documents
# are readable and suitable for academic submission.
# ─────────────────────────────────────────────────────────────────────────────

# Subscript / superscript digit maps for limit notation
_SUBSCRIPT = str.maketrans('0123456789+-=()abcdefghijklmnopqrstuvwxyz',
                            '₀₁₂₃₄₅₆₇₈₉₊₋₌₍₎ₐ' + 'b' + 'cdₑfgₕᵢⱼₖₗₘₙₒₚ' + 'q' + 'ᵣₛₜᵤᵥ' + 'w' + 'ₓ' + 'yz')
_SUPERSCRIPT = str.maketrans('0123456789+-=()abcdefghijklmnopqrstuvwxyz',
                              '⁰¹²³⁴⁵⁶⁷⁸⁹⁺⁻⁼⁽⁾ᵃᵇᶜᵈᵉᶠᵍʰⁱʲᵏˡᵐⁿᵒᵖ' + 'q' + 'ʳˢᵗᵘᵛʷˣʸᶻ')


def _to_sub(s: str) -> str:
    return s.translate(_SUBSCRIPT)


def _to_sup(s: str) -> str:
    return s.translate(_SUPERSCRIPT)


def latex_to_readable(text: str) -> str:
    """
    Convert LaTeX math expressions in *text* to human-readable Unicode equivalents.

    Handles:
    - Display math:  $$...$$ and \\[...\\]
    - Inline math:   $...$ and \\(...\\)
    - \\frac{a}{b}  → (a/b)
    - \\sqrt{x}     → √(x)
    - \\int_{a}^{b} → ∫ₐᵇ
    - \\sum_{i}^{n} → Σᵢⁿ
    - \\prod_{i}^{n}→ Πᵢⁿ
    - \\lim_{x→0}   → lim_{x→0}
    - \\binom{n}{k} → C(n,k)
    - Greek letters, operators, relations, arrows, sets, misc symbols
    - Superscripts ^{...} and subscripts _{...}
    """
    if not text:
        return text

    # ── 1. Structural commands processed first (order matters) ───────────────

    # \frac{numerator}{denominator}  →  (num/den)
    def _frac(m):
        num = latex_to_readable(m.group(1))
        den = latex_to_readable(m.group(2))
        return f'({num}/{den})'
    text = re.sub(r'\\frac\{([^{}]*)\}\{([^{}]*)\}', _frac, text)

    # \sqrt[n]{x}  →  ⁿ√(x)   or   \sqrt{x}  →  √(x)
    def _sqrt(m):
        n = m.group(1)
        inner = latex_to_readable(m.group(2))
        prefix = _to_sup(n) if n else ''
        return f'{prefix}√({inner})'
    text = re.sub(r'\\sqrt(?:\[([^\]]*)\])?\{([^{}]*)\}', _sqrt, text)

    # \binom{n}{k}  →  C(n, k)
    def _binom(m):
        return f'C({latex_to_readable(m.group(1))}, {latex_to_readable(m.group(2))})'
    text = re.sub(r'\\binom\{([^{}]*)\}\{([^{}]*)\}', _binom, text)

    # \int_{lower}^{upper}  →  ∫ₗₒwₑᵣᵘᵖᵖᵉʳ
    # \sum, \prod, \lim with optional limits
    def _big_op(symbol):
        def _replace(m):
            sub = _to_sub(m.group(1)) if m.group(1) else ''
            sup = _to_sup(m.group(2)) if m.group(2) else ''
            return f'{symbol}{sub}{sup} '
        return _replace

    text = re.sub(r'\\int_\{([^{}]*)\}\^\{([^{}]*)\}', _big_op('∫'), text)
    text = re.sub(r'\\int_\{([^{}]*)\}',                lambda m: f'∫{_to_sub(m.group(1))} ', text)
    text = re.sub(r'\\int\^\{([^{}]*)\}',               lambda m: f'∫{_to_sup(m.group(1))} ', text)
    text = re.sub(r'\\int\b',                            '∫', text)

    text = re.sub(r'\\sum_\{([^{}]*)\}\^\{([^{}]*)\}',  _big_op('Σ'), text)
    text = re.sub(r'\\sum_\{([^{}]*)\}',                 lambda m: f'Σ{_to_sub(m.group(1))} ', text)
    text = re.sub(r'\\sum\b',                            'Σ', text)

    text = re.sub(r'\\prod_\{([^{}]*)\}\^\{([^{}]*)\}', _big_op('Π'), text)
    text = re.sub(r'\\prod_\{([^{}]*)\}',                lambda m: f'Π{_to_sub(m.group(1))} ', text)
    text = re.sub(r'\\prod\b',                           'Π', text)

    text = re.sub(r'\\lim_\{([^{}]*)\}',                 lambda m: f'lim({latex_to_readable(m.group(1))}) ', text)
    text = re.sub(r'\\lim\b',                            'lim', text)

    # Superscript: x^{expr} or x^n  (single char)
    text = re.sub(r'\^\{([^{}]+)\}', lambda m: _to_sup(m.group(1)), text)
    text = re.sub(r'\^([0-9a-zA-Z])', lambda m: _to_sup(m.group(1)), text)

    # Subscript: x_{expr} or x_n  (single char)
    text = re.sub(r'_\{([^{}]+)\}', lambda m: _to_sub(m.group(1)), text)
    text = re.sub(r'_([0-9a-zA-Z])', lambda m: _to_sub(m.group(1)), text)

    # \text{...}  →  just the inner text
    text = re.sub(r'\\text\{([^{}]*)\}', r'\1', text)

    # \left( \right)  →  ( )   etc.
    text = re.sub(r'\\left\s*([({[\|])', r'\1', text)
    text = re.sub(r'\\right\s*([)}\]|])', r'\1', text)
    text = re.sub(r'\\left\.', '', text)
    text = re.sub(r'\\right\.', '', text)

    # ── 2. Symbol table (order: longer commands before shorter prefixes) ──────
    _SYMBOLS = [
        # Greek uppercase
        (r'\\Gamma\b',    'Γ'),  (r'\\Delta\b',   'Δ'),  (r'\\Theta\b',   'Θ'),
        (r'\\Lambda\b',   'Λ'),  (r'\\Xi\b',      'Ξ'),  (r'\\Pi\b',      'Π'),
        (r'\\Sigma\b',    'Σ'),  (r'\\Upsilon\b', 'Υ'),  (r'\\Phi\b',     'Φ'),
        (r'\\Psi\b',      'Ψ'),  (r'\\Omega\b',   'Ω'),
        # Greek lowercase
        (r'\\alpha\b',    'α'),  (r'\\beta\b',    'β'),  (r'\\gamma\b',   'γ'),
        (r'\\delta\b',    'δ'),  (r'\\epsilon\b', 'ε'),  (r'\\varepsilon\b','ε'),
        (r'\\zeta\b',     'ζ'),  (r'\\eta\b',     'η'),  (r'\\theta\b',   'θ'),
        (r'\\vartheta\b', 'θ'),  (r'\\iota\b',    'ι'),  (r'\\kappa\b',   'κ'),
        (r'\\lambda\b',   'λ'),  (r'\\mu\b',      'μ'),  (r'\\nu\b',      'ν'),
        (r'\\xi\b',       'ξ'),  (r'\\pi\b',      'π'),  (r'\\varpi\b',   'π'),
        (r'\\rho\b',      'ρ'),  (r'\\varrho\b',  'ρ'),  (r'\\sigma\b',   'σ'),
        (r'\\varsigma\b', 'ς'),  (r'\\tau\b',     'τ'),  (r'\\upsilon\b', 'υ'),
        (r'\\phi\b',      'φ'),  (r'\\varphi\b',  'φ'),  (r'\\chi\b',     'χ'),
        (r'\\psi\b',      'ψ'),  (r'\\omega\b',   'ω'),
        # Operators & relations
        (r'\\times\b',    '×'),  (r'\\div\b',     '÷'),  (r'\\pm\b',      '±'),
        (r'\\mp\b',       '∓'),  (r'\\cdot\b',    '·'),  (r'\\bullet\b',  '•'),
        (r'\\circ\b',     '∘'),  (r'\\star\b',    '★'),  (r'\\ast\b',     '*'),
        (r'\\oplus\b',    '⊕'),  (r'\\otimes\b',  '⊗'),  (r'\\odot\b',    '⊙'),
        (r'\\leq\b',      '≤'),  (r'\\geq\b',     '≥'),  (r'\\neq\b',     '≠'),
        (r'\\le\b',       '≤'),  (r'\\ge\b',      '≥'),  (r'\\ne\b',      '≠'),
        (r'\\ll\b',       '≪'),  (r'\\gg\b',      '≫'),  (r'\\approx\b',  '≈'),
        (r'\\sim\b',      '∼'),  (r'\\simeq\b',   '≃'),  (r'\\cong\b',    '≅'),
        (r'\\equiv\b',    '≡'),  (r'\\propto\b',  '∝'),  (r'\\perp\b',    '⊥'),
        (r'\\parallel\b', '∥'),  (r'\\angle\b',   '∠'),  (r'\\therefore\b','∴'),
        (r'\\because\b',  '∵'),
        # Set / logic
        (r'\\in\b',       '∈'),  (r'\\notin\b',   '∉'),  (r'\\subset\b',  '⊂'),
        (r'\\supset\b',   '⊃'),  (r'\\subseteq\b','⊆'),  (r'\\supseteq\b','⊇'),
        (r'\\cup\b',      '∪'),  (r'\\cap\b',     '∩'),  (r'\\emptyset\b','∅'),
        (r'\\varnothing\b','∅'), (r'\\forall\b',  '∀'),  (r'\\exists\b',  '∃'),
        (r'\\nexists\b',  '∄'),  (r'\\neg\b',     '¬'),  (r'\\lnot\b',    '¬'),
        (r'\\land\b',     '∧'),  (r'\\lor\b',     '∨'),  (r'\\wedge\b',   '∧'),
        (r'\\vee\b',      '∨'),
        # Arrows
        (r'\\rightarrow\b','→'), (r'\\leftarrow\b','←'), (r'\\leftrightarrow\b','↔'),
        (r'\\Rightarrow\b','⇒'), (r'\\Leftarrow\b','⇐'), (r'\\Leftrightarrow\b','⇔'),
        (r'\\to\b',       '→'),  (r'\\gets\b',    '←'),  (r'\\mapsto\b',  '↦'),
        (r'\\uparrow\b',  '↑'),  (r'\\downarrow\b','↓'), (r'\\updownarrow\b','↕'),
        # Misc math
        (r'\\infty\b',    '∞'),  (r'\\partial\b', '∂'),  (r'\\nabla\b',   '∇'),
        (r'\\hbar\b',     'ℏ'),  (r'\\ell\b',     'ℓ'),  (r'\\Re\b',      'ℜ'),
        (r'\\Im\b',       'ℑ'),  (r'\\aleph\b',   'ℵ'),
        (r'\\cdots\b',    '⋯'),  (r'\\ldots\b',   '…'),  (r'\\vdots\b',   '⋮'),
        (r'\\ddots\b',    '⋱'),  (r'\\dots\b',    '…'),
        (r'\\prime\b',    '′'),  (r'\\degree\b',  '°'),
        # Number sets
        (r'\\mathbb\{R\}','ℝ'),  (r'\\mathbb\{Z\}','ℤ'),  (r'\\mathbb\{N\}','ℕ'),
        (r'\\mathbb\{Q\}','ℚ'),  (r'\\mathbb\{C\}','ℂ'),  (r'\\mathbb\{P\}','ℙ'),
        # Trig / functions (remove backslash, keep name)
        (r'\\sin\b',  'sin'),   (r'\\cos\b',  'cos'),   (r'\\tan\b',  'tan'),
        (r'\\sec\b',  'sec'),   (r'\\csc\b',  'csc'),   (r'\\cot\b',  'cot'),
        (r'\\arcsin\b','arcsin'),(r'\\arccos\b','arccos'),(r'\\arctan\b','arctan'),
        (r'\\sinh\b', 'sinh'),  (r'\\cosh\b', 'cosh'),  (r'\\tanh\b', 'tanh'),
        (r'\\log\b',  'log'),   (r'\\ln\b',   'ln'),    (r'\\exp\b',  'exp'),
        (r'\\det\b',  'det'),   (r'\\dim\b',  'dim'),   (r'\\ker\b',  'ker'),
        (r'\\max\b',  'max'),   (r'\\min\b',  'min'),   (r'\\sup\b',  'sup'),
        (r'\\inf\b',  'inf'),   (r'\\arg\b',  'arg'),   (r'\\gcd\b',  'gcd'),
        (r'\\lcm\b',  'lcm'),   (r'\\mod\b',  'mod'),
        # Spacing / formatting commands — strip silently
        (r'\\,', ' '), (r'\\;', ' '), (r'\\:', ' '), (r'\\!', ''),
        (r'\\quad\b', '  '), (r'\\qquad\b', '   '),
        (r'\\\\', '\n'),          # LaTeX line break
        (r'\\newline\b', '\n'),
    ]

    for pattern, replacement in _SYMBOLS:
        text = re.sub(pattern, replacement, text)

    # ── 3. Strip remaining unknown backslash commands ─────────────────────────
    # e.g. \mathbf{x} → x,  \overline{AB} → AB
    text = re.sub(r'\\[a-zA-Z]+\{([^{}]*)\}', r'\1', text)
    # Bare unknown commands with no braces: \foo → (nothing)
    text = re.sub(r'\\[a-zA-Z]+\b', '', text)

    # ── 4. Strip math delimiters ──────────────────────────────────────────────
    # Display math: $$...$$ and \[...\]  → content on its own line
    text = re.sub(r'\$\$(.+?)\$\$', lambda m: f'\n{m.group(1).strip()}\n', text, flags=re.DOTALL)
    text = re.sub(r'\\\[(.+?)\\\]',  lambda m: f'\n{m.group(1).strip()}\n', text, flags=re.DOTALL)
    # Inline math: $...$ and \(...\)
    text = re.sub(r'\$(.+?)\$',      lambda m: m.group(1).strip(), text)
    text = re.sub(r'\\\((.+?)\\\)',  lambda m: m.group(1).strip(), text)

    # ── 5. Strip any leftover lone braces ─────────────────────────────────────
    text = re.sub(r'(?<!\\)[{}]', '', text)

    return text

class AssignmentViewSet(viewsets.ModelViewSet):
    """
    Unified ViewSet for all Assignment operations including creation, 
    detail, AI solving, refinement, and exporting.
    """
    serializer_class = AssignmentSerializer
    permission_classes = [permissions.IsAuthenticated]


    def get_queryset(self):
        from django.db.models import Q
        user = self.request.user
        if user.is_anonymous:
            return Assignment.objects.none()
            
        # Return owned assignments OR assignments shared in a workspace where I'm a member
        return Assignment.objects.filter(
            Q(user=user) | 
            Q(workspace_shares__workspace__memberships__user=user)
        ).distinct()

    def perform_create(self, serializer):
        from .models import AssignmentSource
        
        # Check if any content is provided
        instructions = self.request.data.get('instructions', '').strip()
        file = self.request.FILES.get('file')
        images = self.request.FILES.getlist('image_sources')
        
        if not instructions and not file and not images:
             from rest_framework.exceptions import ValidationError
             raise ValidationError('Initialization failure: No textual, visual, or document sources provided.')

        assignment = serializer.save(user=self.request.user)
        
        # Track main PDF as a source
        if assignment.file:
            AssignmentSource.objects.get_or_create(
                assignment=assignment,
                file=assignment.file,
                file_type='pdf'
            )
            
        # Track additional image sources
        for img in images:
            AssignmentSource.objects.create(
                assignment=assignment,
                file=img,
                file_type='image'
            )

        if assignment.due_date:
            self._create_deadline(assignment)

    def _create_deadline(self, assignment):
        try:
            from planner.models import Deadline
            Deadline.objects.create(
                user=assignment.user,
                title=assignment.title,
                subject=assignment.subject or '',
                due_date=assignment.due_date,
                assignment=assignment,
            )
            # Notify if due within 7 days
            days_until = (assignment.due_date - timezone.now()).days
            if days_until <= 7:
                from users.notifications import notify_deadline_approaching
                notify_deadline_approaching(assignment.user, assignment.title, days_until)
        except Exception as e:
            logger.warning(f'Could not create deadline for assignment {assignment.id}: {e}')

    def get_parsers(self):
        from rest_framework.parsers import MultiPartParser, FormParser, JSONParser
        return [MultiPartParser(), FormParser(), JSONParser()]

    @action(detail=True, methods=['get'])
    def download_intelligence(self, request, pk=None):
        """Export assignment intelligence (Allows shared workspace members)."""
        # Bypass default get_object check for shared access
        from .models import Assignment
        assignment = get_object_or_404(Assignment, pk=pk)
        
        # Security Check: Is Owner OR is a member of a workspace where shared
        can_access = (assignment.user == request.user)
        if not can_access:
             from workspace.models import WorkspaceMember
             can_access = WorkspaceMember.objects.filter(
                 workspace__messages__shared_assignment=assignment,
                 user=request.user
             ).exists()
             
        if not can_access:
            return Response({'error': 'Unauthorized access'}, status=status.HTTP_403_FORBIDDEN)

        export_format = request.query_params.get('format', 'pdf').lower()
        if export_format == 'docx':
            return self._export_docx(assignment)
        elif export_format == 'txt':
            return self._export_txt(assignment)
        return self._export_pdf(assignment)

    @action(detail=True, methods=['post'])
    def share(self, request, pk=None):
        """Transmit assignment intelligence to a target Workspace."""
        assignment = self.get_object()
        workspace_id = request.data.get('workspace_id')
        if not workspace_id:
            return Response({'error': 'Target workspace_id is required.'}, status=status.HTTP_400_BAD_REQUEST)
            
        try:
            from workspace.models import Workspace, WorkspaceMessage
            ws = get_object_or_404(Workspace, id=workspace_id)
            
            # Verify requestor is a member of target workspace
            if not ws.memberships.filter(user=request.user).exists():
                return Response({'error': 'You must be a member of the target workspace to share.'}, status=status.HTTP_403_FORBIDDEN)
            
            # Create Transmission Message
            WorkspaceMessage.objects.create(
                workspace=ws,
                author=request.user,
                content=f"has shared the completed assignment: **{assignment.title}**.",
                shared_assignment=assignment
            )
            
            # Link resources
            ws.resources.add(*assignment.resources.all())
                
            return Response({'message': f'Intelligence transmitted to {ws.name}'})
        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    @action(detail=True, methods=['post'])
    def solve(self, request, pk=None):
        """
        Premium Synthesis Engine: Triggers multi-modal AI processing.
        Supports text instructions, primary PDF, additional PDFs, and images.
        """
        assignment = self.get_object()
        
        # --- Step 1: Build textual context from all PDF sources ---
        # If no text instructions, try to extract from all attached PDFs
        if not assignment.instructions.strip():
            extracted_texts = []
            
            # Try primary file first
            if assignment.file:
                try:
                    from library.pdf_extractor import extract_pdf_text
                    text = extract_pdf_text(assignment.file.path)
                    if text:
                        extracted_texts.append(text[:4000])
                except Exception as e:
                    logger.warning(f'Primary PDF extraction failed: {e}')
            
            # Also extract from any additional PDF sources
            for source in assignment.sources.filter(file_type='pdf'):
                try:
                    from library.pdf_extractor import extract_pdf_text
                    text = extract_pdf_text(source.file.path)
                    if text:
                        extracted_texts.append(text[:2000])
                except Exception as e:
                    logger.warning(f'Source PDF extraction failed ({source.id}): {e}')
            
            if extracted_texts:
                assignment.instructions = '\n\n---\n\n'.join(extracted_texts)
                assignment.save(update_fields=['instructions'])

        # --- Step 2: Validate - must have text OR image sources ---
        has_images = assignment.sources.filter(file_type='image').exists()
        if not assignment.instructions.strip() and not has_images:
            return Response(
                {'error': 'Cannot synthesize: No text instructions, PDF content, or image sources found.'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        # For image-only assignments, inject a placeholder so AI knows to focus on visuals
        if not assignment.instructions.strip() and has_images:
            assignment.instructions = f"[Image-only assignment: '{assignment.title}'. Analyze the attached visual sources and provide a comprehensive academic response to the questions/content depicted in the images.]"
            assignment.save(update_fields=['instructions'])

        # --- Step 3: Mark as processing and call AI ---
        assignment.status = 'processing'
        assignment.save(update_fields=['status'])

        try:
            from ai_assistant.services import AIService
            ai = AIService()
            result = ai.solve_assignment(assignment)

            assignment.ai_response = result.get('response', '')
            assignment.ai_overview = result.get('overview', '')
            if result.get('outline'):
                assignment.ai_outline = result['outline']
            assignment.status = 'completed'
            assignment.save()
            return Response(self.get_serializer(assignment).data)
        except Exception as e:
            logger.error(f'Synthesis engine failure for assignment {assignment.id}: {e}')
            assignment.status = 'error'
            assignment.save(update_fields=['status'])
            return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    @action(detail=True, methods=['post'])
    def refine(self, request, pk=None):
        """Iteratively refine the AI response based on user feedback."""
        assignment = self.get_object()
        prompt = request.data.get('prompt')
        if not prompt:
            return Response({'error': 'Prompt is required.'}, status=status.HTTP_400_BAD_REQUEST)

        try:
            from ai_assistant.services import AIService
            ai = AIService()
            result = ai.refine_assignment(assignment, prompt)
            
            assignment.ai_response = result['response']
            assignment.save()
            return Response(self.get_serializer(assignment).data)
        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    # --- Other Actions (humanize, originality, detect, roadmap, transform, schedule) simplified for brevity in this cleanup, 
    # but I'll ensure they are present in the final file ---

    @action(detail=True, methods=['post'])
    def humanize(self, request, pk=None):
        """Engage Vanish Protocol to bypass AI detection."""
        assignment = self.get_object()
        try:
            from ai_assistant.services import AIService
            ai = AIService()
            result = ai.humanize_assignment(assignment)
            assignment.ai_response = result['response']
            assignment.ai_overview = result['overview']
            assignment.chat_history = result.get('chat_history', [])
            assignment.save()
            return Response(self.get_serializer(assignment).data)
        except Exception as e: return Response({'error': str(e)}, status=500)

    @action(detail=True, methods=['post'])
    def originality(self, request, pk=None):
        """Engage Originality Shield to remove plagiarism markers."""
        assignment = self.get_object()
        try:
            from ai_assistant.services import AIService
            ai = AIService()
            result = ai.remove_plagiarism(assignment)
            assignment.ai_response = result['response']
            assignment.ai_overview = result['overview']
            assignment.chat_history = result.get('chat_history', [])
            assignment.save()
            return Response(self.get_serializer(assignment).data)
        except Exception as e: return Response({'error': str(e)}, status=500)

    @action(detail=True, methods=['post'])
    def detect(self, request, pk=None):
        """Perform a deep-fidelity linguistic audit for AI and plagiarism."""
        assignment = self.get_object()
        try:
            from ai_assistant.services import AIService
            ai = AIService()
            result = ai.detect_assignment(assignment)
            return Response(result)
        except Exception as e: return Response({'error': str(e)}, status=500)

    @action(detail=True, methods=['post'])
    def roadmap(self, request, pk=None):
        """Generate a study roadmap for the assignment content."""
        assignment = self.get_object()
        try:
            from ai_assistant.services import AIService
            ai = AIService()
            roadmap = ai.generate_assignment_roadmap(assignment)
            return Response({'message': f'Generated {len(roadmap)} milestones.'})
        except Exception as e: return Response({'error': str(e)}, status=500)

    # --- Export Helper Methods ---

    def _export_txt(self, assignment):
        content = f"{assignment.title}\n\n{latex_to_readable(assignment.ai_response or '')}"
        response = HttpResponse(content, content_type='text/plain')
        response['Content-Disposition'] = f'attachment; filename="{assignment.title}.txt"'
        return response

    def _export_pdf(self, assignment):
        """Premium Synthesis Engine: Converts Markdown into a high-fidelity academic document."""
        from reportlab.lib.pagesizes import A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_LEFT
        from reportlab.lib.colors import HexColor
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer, 
            pagesize=A4,
            rightMargin=72, leftMargin=72,
            topMargin=72, bottomMargin=72
        )
        
        styles = getSampleStyleSheet()
        
        # --- Premium Typographic Tokens ---
        title_style = ParagraphStyle(
            'PremiumTitle',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=30,
            textColor=HexColor('#0f172a'),
            fontName='Helvetica-Bold'
        )
        
        h2_style = ParagraphStyle(
            'PremiumH2', parent=styles['Heading2'], fontSize=16, spaceBefore=22, spaceAfter=12,
            textColor=HexColor('#1e293b'), fontName='Helvetica-Bold', leading=20
        )

        h3_style = ParagraphStyle(
            'PremiumH3', parent=styles['Heading3'], fontSize=13, spaceBefore=16, spaceAfter=8,
            textColor=HexColor('#334155'), fontName='Helvetica-Bold'
        )
        
        body_style = ParagraphStyle(
            'PremiumBody', parent=styles['Normal'], fontSize=11, leading=16,
            alignment=TA_LEFT, textColor=HexColor('#475569'), spaceAfter=10
        )

        list_style = ParagraphStyle(
            'PremiumList', parent=body_style, leftIndent=20, bulletIndent=10, spaceAfter=6
        )

        # --- Enhanced Markdown Parser (Handles unformatted/single-line blocks) ---
        def parse_content(text):
            if not text: return []
            
            # Pre-processing: If there are no actual newlines but there are markdown markers, force some breaks
            if '\n' not in text and '###' in text:
                text = text.replace('###', '\n###').replace('##', '\n##').replace('#', '\n#')

            lines = text.split('\n')
            story = []
            
            for line in lines:
                line = line.strip()
                if not line:
                    story.append(Spacer(1, 6))
                    continue
                
                # Inline Bold/Italic
                line = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', line)
                line = re.sub(r'\*(.*?)\*', r'<i>\1</i>', line)
                
                # Headers
                if line.startswith('### '):
                    story.append(Paragraph(line[4:], h3_style))
                elif line.startswith('## '):
                    story.append(Paragraph(line[3:], h2_style))
                elif line.startswith('# '):
                    story.append(Paragraph(line[2:], h2_style))
                # Lists
                elif line.startswith('- ') or line.startswith('* '):
                    story.append(Paragraph(f"• {line[2:]}", list_style))
                elif re.match(r'^\d+\.\s', line):
                    story.append(Paragraph(line, list_style))
                else:
                    story.append(Paragraph(line, body_style))
            
            return story

        story = [Paragraph(assignment.title, title_style), Spacer(1, 2)]
        
        story.append(Paragraph(
            f"Subject: {assignment.subject or 'General Integration'}", 
            ParagraphStyle('Meta', parent=body_style, fontSize=9, textColor=HexColor('#64748b'))
        ))
        story.append(Spacer(1, 24))
        
        story.extend(parse_content(latex_to_readable(assignment.ai_response or '')))
        
        def add_footer(canvas, doc):
            canvas.saveState()
            canvas.setFont('Helvetica', 8)
            canvas.setStrokeColor(HexColor('#e2e8f0'))
            canvas.line(72, 50, 523, 50)
            canvas.drawRightString(523, 40, f"Page {doc.page}")
            canvas.restoreState()

        doc.build(story, onFirstPage=add_footer, onLaterPages=add_footer)
        
        buffer.seek(0)
        data = buffer.read()
        response = HttpResponse(data, content_type='application/pdf')
        filename = self._safe_filename(assignment.title)
        response['Content-Disposition'] = f'attachment; filename="{filename}.pdf"'
        return response

    def _export_docx(self, assignment):
        from docx import Document
        doc = Document()
        doc.add_heading(assignment.title, 0)
        doc.add_paragraph(latex_to_readable(assignment.ai_response or ''))
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        response = HttpResponse(buffer.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        filename = self._safe_filename(assignment.title)
        response['Content-Disposition'] = f'attachment; filename="{filename}.docx"'
        return response

    def _safe_filename(self, title):
        if not title: return "assignment"
        return re.sub(r'[^\w\s-]', '', title).strip().replace(' ', '_')[:50]


# ─────────────────────────────────────────────────────────────────────────────
# STANDALONE EXPORT VIEW
# Using a plain @api_view instead of ViewSet.as_view() to guarantee the URL
# always resolves. ViewSet.as_view() outside the router can silently 404.
# ─────────────────────────────────────────────────────────────────────────────

from rest_framework.decorators import api_view, permission_classes as pc
from rest_framework.permissions import IsAuthenticated

def _safe_filename_util(title):
    if not title: return "assignment"
    return re.sub(r'[^\w\s-]', '', title).strip().replace(' ', '_')[:50]


@api_view(['GET'])
@pc([IsAuthenticated])
def export_assignment(request, pk):
    """Standalone export view — guaranteed URL resolution, no ViewSet routing."""
    assignment = get_object_or_404(Assignment, pk=pk)

    # Auth: must be owner OR workspace member with access
    if assignment.user != request.user:
        try:
            from workspace.models import WorkspaceMember
            has_access = WorkspaceMember.objects.filter(
                workspace__messages__shared_assignment=assignment,
                user=request.user
            ).exists()
        except Exception:
            has_access = False
        if not has_access:
            return Response({'error': 'Unauthorized'}, status=403)

    export_format = request.query_params.get('format', 'pdf').lower()

    if export_format == 'txt':
        content = f"{assignment.title}\n\n{latex_to_readable(assignment.ai_response or '')}"
        resp = HttpResponse(content, content_type='text/plain')
        resp['Content-Disposition'] = f'attachment; filename="{_safe_filename_util(assignment.title)}.txt"'
        return resp

    if export_format == 'docx':
        try:
            from docx import Document
            doc = Document()
            doc.add_heading(assignment.title, 0)
            doc.add_paragraph(latex_to_readable(assignment.ai_response or ''))
            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)
            resp = HttpResponse(buf.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            resp['Content-Disposition'] = f'attachment; filename="{_safe_filename_util(assignment.title)}.docx"'
            return resp
        except ImportError:
            return Response({'error': 'Word export not available. python-docx not installed.'}, status=500)
        except Exception as e:
            return Response({'error': str(e)}, status=500)

    # Default: PDF
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_LEFT
        from reportlab.lib.colors import HexColor

        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=72)
        styles = getSampleStyleSheet()

        title_style = ParagraphStyle('T', parent=styles['Heading1'], fontSize=22, spaceAfter=24,
                                     textColor=HexColor('#0f172a'), fontName='Helvetica-Bold')
        h2_style = ParagraphStyle('H2', parent=styles['Heading2'], fontSize=15, spaceBefore=18,
                                  spaceAfter=10, textColor=HexColor('#1e293b'), fontName='Helvetica-Bold')
        h3_style = ParagraphStyle('H3', parent=styles['Heading3'], fontSize=12, spaceBefore=12,
                                  spaceAfter=6, textColor=HexColor('#334155'), fontName='Helvetica-Bold')
        body_style = ParagraphStyle('B', parent=styles['Normal'], fontSize=11, leading=17,
                                    alignment=TA_LEFT, textColor=HexColor('#374151'), spaceAfter=8)
        list_style = ParagraphStyle('L', parent=body_style, leftIndent=18, spaceAfter=5)

        def parse_md(text):
            if not text: return []
            story = []
            for line in text.split('\n'):
                line = line.strip()
                if not line:
                    story.append(Spacer(1, 6))
                    continue
                line = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', line)
                line = re.sub(r'\*(.*?)\*', r'<i>\1</i>', line)
                if line.startswith('### '):
                    story.append(Paragraph(line[4:], h3_style))
                elif line.startswith('## '):
                    story.append(Paragraph(line[3:], h2_style))
                elif line.startswith('# '):
                    story.append(Paragraph(line[2:], h2_style))
                elif line.startswith(('- ', '* ')):
                    story.append(Paragraph(f"• {line[2:]}", list_style))
                elif re.match(r'^\d+\.\s', line):
                    story.append(Paragraph(line, list_style))
                else:
                    story.append(Paragraph(line, body_style))
            return story

        story = [Paragraph(assignment.title, title_style), Spacer(1, 4)]
        meta = ParagraphStyle('M', parent=body_style, fontSize=9, textColor=HexColor('#64748b'))
        story.append(Paragraph(f"Subject: {assignment.subject or 'General'}", meta))
        story.append(Spacer(1, 20))
        story.extend(parse_md(latex_to_readable(assignment.ai_response or '')))

        def footer(canvas, doc):
            canvas.saveState()
            canvas.setFont('Helvetica', 8)
            canvas.setStrokeColor(HexColor('#e2e8f0'))
            canvas.line(72, 50, 523, 50)
            canvas.drawRightString(523, 40, f"Page {doc.page}")
            canvas.restoreState()

        doc.build(story, onFirstPage=footer, onLaterPages=footer)
        buf.seek(0)
        resp = HttpResponse(buf.read(), content_type='application/pdf')
        resp['Content-Disposition'] = f'attachment; filename="{_safe_filename_util(assignment.title)}.pdf"'
        return resp

    except ImportError:
        return Response({'error': 'PDF export not available. reportlab not installed.'}, status=500)
    except Exception as e:
        logger.error(f'PDF export failed for assignment {pk}: {e}')
        return Response({'error': str(e)}, status=500)
