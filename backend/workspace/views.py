import io
import logging
import re
from rest_framework import generics, permissions, status
from rest_framework.response import Response
from rest_framework.views import APIView
from django.shortcuts import get_object_or_404
from django.http import HttpResponse
from django.utils import timezone
from .models import Workspace, WorkspaceMember, WorkspaceBlock, WorkspaceMessage, WorkspaceTask, DocumentVersion, WorkspaceFile
from .serializers import (
    WorkspaceSerializer, WorkspaceMemberSerializer, WorkspaceBlockSerializer,
    WorkspaceMessageSerializer, WorkspaceTaskSerializer, DocumentVersionSerializer,
    WorkspaceFileSerializer
)

logger = logging.getLogger('flowstate')


def _get_workspace(pk, user):
    ws = get_object_or_404(Workspace, pk=pk)
    if not ws.memberships.filter(user=user).exists():
        from rest_framework.exceptions import PermissionDenied
        raise PermissionDenied('You are not a member of this workspace.')
    return ws


class WorkspaceListCreateView(generics.ListCreateAPIView):
    permission_classes = [permissions.IsAuthenticated]
    serializer_class = WorkspaceSerializer

    def get_queryset(self):
        return Workspace.objects.filter(memberships__user=self.request.user, is_active=True)

    def perform_create(self, serializer):
        ws = serializer.save(owner=self.request.user)
        WorkspaceMember.objects.create(workspace=ws, user=self.request.user, role='owner')
        # Create initial welcome block
        WorkspaceBlock.objects.create(
            workspace=ws,
            block_type='text',
            content='# Welcome to your new workspace!\n\nFlowAI is ready to help you collaborate.',
            order=0
        )

    def get_serializer_context(self):
        return {'request': self.request}


class WorkspaceDetailView(generics.RetrieveUpdateDestroyAPIView):
    permission_classes = [permissions.IsAuthenticated]
    serializer_class = WorkspaceSerializer

    def get_queryset(self):
        return Workspace.objects.filter(memberships__user=self.request.user)

    def get_serializer_context(self):
        return {'request': self.request}


class JoinWorkspaceView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def post(self, request):
        code = request.data.get('invite_code', '').strip().upper()
        if not code:
            return Response({'error': 'Invite code required.'}, status=status.HTTP_400_BAD_REQUEST)
        ws = get_object_or_404(Workspace, invite_code=code, is_active=True)
        member, created = WorkspaceMember.objects.get_or_create(
            workspace=ws, user=request.user,
            defaults={'role': 'editor'}
        )
        if not created:
            return Response({'detail': 'Already a member.'})
        return Response(WorkspaceSerializer(ws, context={'request': request}).data, status=status.HTTP_201_CREATED)


class WorkspaceMembersView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request, pk):
        ws = _get_workspace(pk, request.user)
        # Update last_seen
        WorkspaceMember.objects.filter(workspace=ws, user=request.user).update(last_seen=timezone.now())
        members = ws.memberships.select_related('user').all()
        return Response(WorkspaceMemberSerializer(members, many=True).data)


class WorkspaceBlocksView(APIView):
    """View to list, create, and update blocks in a workspace."""
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request, workspace_id):
        ws = _get_workspace(workspace_id, request.user)
        blocks = ws.blocks.all().order_by('order', 'created_at')
        return Response(WorkspaceBlockSerializer(blocks, many=True).data)

    def post(self, request, workspace_id):
        ws = _get_workspace(workspace_id, request.user)
        serializer = WorkspaceBlockSerializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        
        # Auto-calculate order if not provided
        order = request.data.get('order')
        if order is None:
            last_block = ws.blocks.last()
            order = (last_block.order + 1) if last_block else 0
            
        serializer.save(workspace=ws, order=order, last_edited_by=request.user)
        return Response(serializer.data, status=status.HTTP_201_CREATED)

    def patch(self, request, workspace_id):
        """Bulk update block order or specific block content."""
        ws = _get_workspace(workspace_id, request.user)
        block_id = request.data.get('block_id')
        if not block_id:
            # Handle reordering if a list of IDs is sent
            block_ids = request.data.get('order_list', [])
            for i, bid in enumerate(block_ids):
                WorkspaceBlock.objects.filter(id=bid, workspace=ws).update(order=i)
            return Response({'status': 'reordered'})

        block = get_object_or_404(WorkspaceBlock, id=block_id, workspace=ws)
        serializer = WorkspaceBlockSerializer(block, data=request.data, partial=True)
        serializer.is_valid(raise_exception=True)
        serializer.save(last_edited_by=request.user)
        
        # Trigger version snapshot occasionally (logic handled in consumer later for real-time)
        return Response(serializer.data)

    def delete(self, request, pk):
        ws = _get_workspace(pk, request.user)
        block_id = request.data.get('block_id')
        block = get_object_or_404(WorkspaceBlock, id=block_id, workspace=ws)
        block.delete()
        return Response(status=status.HTTP_204_NO_CONTENT)


class DocumentVersionsView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request, workspace_id):
        ws = _get_workspace(workspace_id, request.user)
        versions = ws.versions.all().order_by('-created_at')
        return Response(DocumentVersionSerializer(versions, many=True).data)

    def post(self, request, workspace_id):
        """Create a new manual snapshot of the workspace."""
        ws = _get_workspace(workspace_id, request.user)
        blocks = ws.blocks.all().order_by('order')
        snapshot = [
            {
                'block_type': b.block_type,
                'content': b.content,
                'metadata': b.metadata,
                'order': b.order
            }
            for b in blocks
        ]
        
        last_v = ws.versions.order_by('-version').first()
        new_v_num = (last_v.version + 1) if last_v else 1
        
        v = DocumentVersion.objects.create(
            workspace=ws,
            blocks_snapshot=snapshot,
            saved_by=request.user,
            version=new_v_num
        )
        return Response(DocumentVersionSerializer(v).data, status=status.HTTP_201_CREATED)

    def put(self, request, workspace_id):
        """Restore a version (restores all blocks from snapshot)."""
        ws = _get_workspace(workspace_id, request.user)
        version_id = request.data.get('version_id')
        v = get_object_or_404(DocumentVersion, id=version_id, workspace=ws)
        
        # Clear current blocks and restore from snapshot
        ws.blocks.all().delete()
        for b_data in v.blocks_snapshot:
            WorkspaceBlock.objects.create(
                workspace=ws,
                block_type=b_data.get('block_type', 'text'),
                content=b_data.get('content', ''),
                metadata=b_data.get('metadata', {}),
                order=b_data.get('order', 0)
            )
        return Response({'status': 'restored', 'version': v.version})


class MessagesView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request, workspace_id):
        ws = _get_workspace(workspace_id, request.user)
        msgs = ws.messages.select_related('author').order_by('-created_at')[:50]
        return Response(WorkspaceMessageSerializer(list(reversed(msgs)), many=True).data)

    def post(self, request, workspace_id):
        ws = _get_workspace(workspace_id, request.user)
        content = request.data.get('content', '').strip()
        if not content:
            return Response({'error': 'Content required.'}, status=status.HTTP_400_BAD_REQUEST)

        msg = WorkspaceMessage.objects.create(workspace=ws, author=request.user, content=content)
        response_data = {'user_message': WorkspaceMessageSerializer(msg).data, 'ai_message': None}

        # If message starts with @FlowAI, trigger AI response and return it immediately
        if content.lower().startswith('@flowai') or content.lower().startswith('@ai'):
            question = re.sub(r'^@(flowai|ai)\s*', '', content, flags=re.IGNORECASE).strip()
            if question:
                ai_reply = _get_ai_response(ws, question, request.user)
                ai_msg = WorkspaceMessage.objects.create(workspace=ws, content=ai_reply, is_ai=True)
                response_data['ai_message'] = WorkspaceMessageSerializer(ai_msg).data

        return Response(response_data, status=status.HTTP_201_CREATED)


class TasksView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request, workspace_id):
        ws = _get_workspace(workspace_id, request.user)
        tasks = ws.tasks.select_related('assigned_to', 'created_by').all()
        return Response(WorkspaceTaskSerializer(tasks, many=True).data)

    def post(self, request, workspace_id):
        ws = _get_workspace(workspace_id, request.user)
        serializer = WorkspaceTaskSerializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        serializer.save(workspace=ws, created_by=request.user)
        return Response(serializer.data, status=status.HTTP_201_CREATED)


class TaskDetailView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def patch(self, request, workspace_id, task_id):
        ws = _get_workspace(workspace_id, request.user)
        task = get_object_or_404(WorkspaceTask, id=task_id, workspace=ws)
        serializer = WorkspaceTaskSerializer(task, data=request.data, partial=True)
        serializer.is_valid(raise_exception=True)
        serializer.save()
        return Response(serializer.data)

    def delete(self, request, workspace_id, task_id):
        ws = _get_workspace(workspace_id, request.user)
        task = get_object_or_404(WorkspaceTask, id=task_id, workspace=ws)
        task.delete()
        return Response(status=status.HTTP_204_NO_CONTENT)


class FilesView(APIView):
    """Upload files directly to a workspace or link library resources."""
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request, pk):
        ws = _get_workspace(pk, request.user)
        files = ws.files.select_related('uploaded_by').all()
        linked = ws.resources.all()
        return Response({
            'uploaded': WorkspaceFileSerializer(files, many=True, context={'request': request}).data,
            'linked_resources': [
                {'id': r.id, 'title': r.title, 'type': r.resource_type, 'subject': r.subject}
                for r in linked
            ],
        })

    def post(self, request, pk):
        ws = _get_workspace(pk, request.user)
        uploaded_file = request.FILES.get('file')
        resource_id = request.data.get('resource_id')

        # Link an existing library resource
        if resource_id:
            from library.models import Resource
            resource = get_object_or_404(Resource, id=resource_id, owner=request.user)
            ws.resources.add(resource)
            return Response({'linked': True, 'resource_id': resource_id})

        # Upload a new file
        if not uploaded_file:
            return Response({'error': 'No file provided.'}, status=status.HTTP_400_BAD_REQUEST)

        import os
        name = uploaded_file.name
        ext = os.path.splitext(name)[1].lower()
        if ext not in ['.pdf', '.doc', '.docx', '.txt', '.png', '.jpg', '.jpeg']:
            return Response({'error': 'Unsupported file type.'}, status=status.HTTP_400_BAD_REQUEST)
        if uploaded_file.size > 20 * 1024 * 1024:
            return Response({'error': 'File too large. Max 20MB.'}, status=status.HTTP_400_BAD_REQUEST)

        ws_file = WorkspaceFile.objects.create(
            workspace=ws,
            uploaded_by=request.user,
            file=uploaded_file,
            name=name,
            file_size=uploaded_file.size,
        )

        # Extract text from PDF for AI context
        if ext == '.pdf':
            try:
                from library.pdf_extractor import extract_pdf_content
                file_bytes = ws_file.file.read()
                pdf_data = extract_pdf_content(file_bytes=file_bytes)
                if pdf_data['text']:
                    ws_file.extracted_text = pdf_data['text'][:20000]
                    ws_file.save(update_fields=['extracted_text'])
            except Exception as e:
                logger.warning(f'Could not extract text from workspace file: {e}')

        return Response(
            WorkspaceFileSerializer(ws_file, context={'request': request}).data,
            status=status.HTTP_201_CREATED
        )

    def delete(self, request, pk):
        ws = _get_workspace(pk, request.user)
        file_id = request.query_params.get('file_id')
        resource_id = request.query_params.get('resource_id')

        if file_id:
            f = get_object_or_404(WorkspaceFile, id=file_id, workspace=ws)
            f.file.delete(save=False)
            f.delete()
            return Response(status=status.HTTP_204_NO_CONTENT)

        if resource_id:
            from library.models import Resource
            try:
                ws.resources.remove(int(resource_id))
            except Exception:
                pass
            return Response(status=status.HTTP_204_NO_CONTENT)

        return Response({'error': 'Provide file_id or resource_id.'}, status=status.HTTP_400_BAD_REQUEST)


class AIAssistView(APIView):
    """FlowAI workspace assistant — full context access."""
    permission_classes = [permissions.IsAuthenticated]

    def post(self, request, workspace_id):
        ws = _get_workspace(workspace_id, request.user)
        action = request.data.get('action', 'ask')  # ask | generate_outline | improve | expand | simplify | review | generate_slides | write_section
        text = request.data.get('text', '')
        question = request.data.get('question', '')

        try:
            if action == 'generate_outline':
                result = _ai_generate_outline(ws)
            elif action == 'improve':
                result = _ai_improve_text(ws, text)
            elif action == 'expand':
                result = _ai_expand_text(ws, text)
            elif action == 'simplify':
                result = _ai_simplify_text(ws, text)
            elif action == 'review':
                result = _ai_review_doc(ws)
            elif action == 'write_section':
                result = _ai_write_section(ws, text)
            elif action == 'generate_slides':
                result = _ai_generate_slides_outline(ws)
            elif action == 'generate_flashcards':
                block_id = request.data.get('block_id')
                block = get_object_or_404(WorkspaceBlock, id=block_id, workspace=ws)
                result = _ai_generate_flashcards(ws, block.content, request.user)
            else:
                result = _get_ai_response(ws, question or text, request.user)

            # Save AI response as chat message
            WorkspaceMessage.objects.create(workspace=ws, content=result, is_ai=True)
            return Response({'result': result})
        except Exception as e:
            logger.error(f'Workspace AI error: {e}')
            return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


class ExportView(APIView):
    """Export workspace document as PDF, DOCX, PPTX, or TXT."""
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request, pk):
        ws = _get_workspace(pk, request.user)
        fmt = request.query_params.get('format', 'pdf').lower()
        
        # Concatenate block contents for export
        blocks = ws.blocks.all().order_by('order', 'created_at')
        full_content = '\n\n'.join([b.content for b in blocks if b.content])

        if fmt == 'pdf':
            return _export_pdf(ws, full_content)
        elif fmt == 'docx':
            return _export_docx(ws, full_content)
        elif fmt == 'pptx':
            return _export_pptx(ws, full_content)
        elif fmt == 'txt':
            return _export_txt(ws, full_content)
        return Response({'error': 'Unsupported format.'}, status=status.HTTP_400_BAD_REQUEST)


# ─── AI HELPERS (Updated for Blocks) ──────────────────────────────────────────

def _build_workspace_context(ws) -> str:
    """Build full workspace context for FlowAI from Blocks and Resources."""
    parts = [f"Workspace: {ws.name}", f"Subject: {ws.subject or 'General'}"]

    if ws.description:
        parts.append(f"Description: {ws.description}")

    if ws.assignment:
        parts.append(f"Assignment: {ws.assignment.title}\nInstructions: {ws.assignment.instructions[:1000]}")

    # Concatenate blocks for context
    blocks = ws.blocks.all()[:15]
    if blocks:
        block_text = '\n'.join([f"[{b.block_type}] {b.content[:1000]}" for b in blocks])
        parts.append(f"Workspace Blocks:\n{block_text}")

    # Recent chat
    recent_msgs = ws.messages.order_by('-created_at')[:8]
    if recent_msgs:
        chat = '\n'.join([f"{'FlowAI' if m.is_ai else (m.author.username if m.author else 'User')}: {m.content}" for m in reversed(recent_msgs)])
        parts.append(f"Recent chat:\n{chat}")

    # Resources (library)
    for r in ws.resources.all()[:3]:
        from ai_assistant.services import AIService
        ai = AIService()
        ctx = ai._get_resource_context(r)
        if ctx:
            parts.append(f"Reference Resource '{r.title}':\n{ctx[:2000]}")

    return '\n\n'.join(parts)


def _get_ai_response(ws, question: str, user) -> str:
    from ai_assistant.services import AIService, FLOWAI_SYSTEM_PROMPT
    ai = AIService()
    context = _build_workspace_context(ws)
    system = (
        f"{FLOWAI_SYSTEM_PROMPT}\n\n"
        f"You are FlowAI inside a collaborative workspace. You have FULL ACCESS to everything in this workspace.\n\n"
        f"WORKSPACE CONTEXT:\n{context}"
    )
    return ai.chat([{'role': 'system', 'content': system}, {'role': 'user', 'content': question}])


def _ai_generate_outline(ws) -> str:
    from ai_assistant.services import AIService, FLOWAI_SYSTEM_PROMPT
    ai = AIService()
    context = _build_workspace_context(ws)
    prompt = (
        f"Based on this workspace context, generate a detailed document outline with sections and subsections.\n\n"
        f"{context}\n\n"
        "Return a well-structured markdown outline that the team can use as their document skeleton. "
        "Include an introduction, main sections based on the assignment/topic, and a conclusion."
    )
    return ai.chat([{'role': 'user', 'content': prompt}])


def _ai_improve_text(ws, text: str) -> str:
    from ai_assistant.services import AIService
    ai = AIService()
    prompt = (
        f"Improve this text to be more academic, clear, and well-structured.\n\n"
        f"Text:\n{text}\n\n"
        "Return only the improved version, no explanation."
    )
    return ai.chat([{'role': 'user', 'content': prompt}])


def _ai_expand_text(ws, text: str) -> str:
    from ai_assistant.services import AIService
    ai = AIService()
    context = _build_workspace_context(ws)
    prompt = (
        f"Expand this text with more detail, using the workspace context where relevant.\n\n"
        f"Context Snippet: {context[:1000]}\n\nText:\n{text}\n\nReturn only the expanded version."
    )
    return ai.chat([{'role': 'user', 'content': prompt}])


def _ai_simplify_text(ws, text: str) -> str:
    from ai_assistant.services import AIService
    ai = AIService()
    prompt = f"Simplify this text while keeping all key information:\n\n{text}\n\nReturn only the simplified version."
    return ai.chat([{'role': 'user', 'content': prompt}])


def _ai_review_doc(ws) -> str:
    from ai_assistant.services import AIService
    ai = AIService()
    blocks = ws.blocks.all().order_by('order')
    content = '\n'.join([b.content for b in blocks if b.content])
    if not content.strip():
        return "The workspace is empty. Add some blocks and I'll review it for you!"
        
    prompt = (
        f"Review this collective work for the workspace '{ws.name}'.\n\n"
        f"Content:\n{content[:6000]}\n\n"
        "Provide structured feedback on quality, gaps, and specific suggestions to improve."
    )
    return ai.chat([{'role': 'user', 'content': prompt}])


def _ai_write_section(ws, section_title: str) -> str:
    from ai_assistant.services import AIService
    ai = AIService()
    context = _build_workspace_context(ws)
    prompt = (
        f"Write a complete section titled '{section_title}' based on the workspace context.\n\n"
        f"Context:\n{context}\n\n"
        "Write in academic style with proper paragraphs. Use markdown formatting."
    )
    return ai.chat([{'role': 'user', 'content': prompt}])


def _ai_generate_slides_outline(ws) -> str:
    from ai_assistant.services import AIService
    ai = AIService()
    context = _build_workspace_context(ws)
    prompt = (
        f"Create a PowerPoint presentation outline for this workspace.\n\n"
        f"Context:\n{context[:3000]}\n\n"
        "Generate a slide-by-slide outline (Title, Agenda, 5-8 Content slides, Conclusion)."
    )
    return ai.chat([{'role': 'user', 'content': prompt}])


# ─── EXPORT HELPERS (Updated) ────────────────────────────────────────────────

def _safe_filename(name):
    return re.sub(r'[^\w\s-]', '', name).strip().replace(' ', '_')[:50]


def _export_txt(ws, content):
    header = f"{ws.name}\n{'='*len(ws.name)}\n\n"
    resp = HttpResponse(header + content, content_type='text/plain; charset=utf-8')
    resp['Content-Disposition'] = f'attachment; filename="{_safe_filename(ws.name)}.txt"'
    return resp


def _export_pdf(ws, content):
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib.units import cm
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

        buffer = io.BytesIO()
        pdf = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=2*cm, leftMargin=2*cm, topMargin=2*cm, bottomMargin=2*cm)
        styles = getSampleStyleSheet()
        story = [Paragraph(ws.name, styles['Title'])]
        if ws.subject:
            story.append(Paragraph(ws.subject, styles['Normal']))
        story.append(Spacer(1, 0.5*cm))

        for line in content.split('\n'):
            line = line.strip()
            if not line:
                story.append(Spacer(1, 0.2*cm))
            elif line.startswith('# '):
                story.append(Paragraph(line[2:], styles['Heading1']))
            elif line.startswith('## '):
                story.append(Paragraph(line[3:], styles['Heading2']))
            elif line.startswith('### '):
                story.append(Paragraph(line[4:], styles['Heading3']))
            elif line.startswith('- ') or line.startswith('* '):
                story.append(Paragraph(f'• {line[2:]}', styles['Normal']))
            else:
                line = re.sub(r'\*\*(.+?)\*\*', r'\1', line)
                story.append(Paragraph(line, styles['Normal']))

        pdf.build(story)
        buffer.seek(0)
        resp = HttpResponse(buffer.read(), content_type='application/pdf')
        resp['Content-Disposition'] = f'attachment; filename="{_safe_filename(ws.name)}.pdf"'
        return resp
    except Exception:
        return _export_txt(ws, content)


def _export_docx(ws, content):
    try:
        from docx import Document as DocxDoc
        d = DocxDoc()
        d.add_heading(ws.name, 0)
        if ws.subject:
            d.add_paragraph(ws.subject)
        for line in content.split('\n'):
            line = line.strip()
            if not line: continue
            if line.startswith('# '): d.add_heading(line[2:], level=1)
            elif line.startswith('## '): d.add_heading(line[3:], level=2)
            elif line.startswith('### '): d.add_heading(line[4:], level=3)
            elif line.startswith('- ') or line.startswith('* '): d.add_paragraph(line[2:], style='List Bullet')
            else:
                line = re.sub(r'\*\*(.+?)\*\*', r'\1', line)
                d.add_paragraph(line)
        buf = io.BytesIO()
        d.save(buf)
        buf.seek(0)
        resp = HttpResponse(buf.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        resp['Content-Disposition'] = f'attachment; filename="{_safe_filename(ws.name)}.docx"'
        return resp
    except Exception:
        return _export_txt(ws, content)

def _ai_generate_flashcards(ws, block_text: str, user) -> str:
    from ai_assistant.services import AIService
    from library.models import Deck, Flashcard
    import json

    ai = AIService()
    prompt = (
        f"Based on the following content from my {ws.name} workspace, generate a set of essential study flashcards.\n\n"
        f"Content:\n{block_text}\n\n"
        "Return ONLY a JSON array of objects with 'question' and 'answer' keys. "
        "Example: [{\"question\": \"...\", \"answer\": \"...\"}, ...]"
    )
    
    response_text = ai.chat([{'role': 'user', 'content': prompt}])
    
    try:
        # Clean response if it contains markdown code blocks
        clean_json = response_text.replace('```json', '').replace('```', '').strip()
        cards_data = json.loads(clean_json)
        
        # Create a workspace deck if it doesn't exist
        deck, _ = Deck.objects.get_or_create(
            owner=user,
            title=f"Workspace: {ws.name}",
            defaults={'subject': ws.subject or 'General'}
        )
        
        created_count = 0
        for card in cards_data:
            if 'question' in card and 'answer' in card:
                Flashcard.objects.create(
                    deck=deck,
                    owner=user,
                    question=card['question'],
                    answer=card['answer'],
                    subject=ws.subject or 'General',
                    resource=ws.resources.first() # Link to first workspace resource if any
                )
                created_count += 1
        
        return f"✨ Success! Generated {created_count} flashcards and added them to your '{deck.title}' deck."
    except Exception as e:
        logger.error(f"Flashcard generation failed: {e}")
        return f"AI returned: {response_text}\n\n(Error parsing cards: {str(e)})"

def _export_pptx(ws, content):
    # Basic fall-back to text for PPTX for now as it needs structured slide content
    return _export_txt(ws, content)
